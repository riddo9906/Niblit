#!/usr/bin/env python3
"""
modules/universal_file_manager.py — Niblit Universal File Manager

Provides a single, extensible entry point for reading, writing, editing,
and executing files of virtually any type.  Each file type is handled by
a lightweight *handler* registered in a central registry; unknown types
fall back to raw UTF-8 text I/O.

Supported formats (all optional-import — degrade gracefully if lib missing)
---------------------------------------------------------------------------
Text / data  : plain text, CSV, JSON, YAML, TOML, INI/CFG
Office / docs: PDF (pdfplumber / PyPDF2), DOCX (python-docx), XLSX/ODS (openpyxl)
Images       : PNG, JPEG, GIF, BMP, TIFF (Pillow)
Audio        : WAV, FLAC, MP3, OGG (soundfile / pydub)
Archives     : ZIP, TAR (stdlib)
Disk images  : ISO (pycdlib)
Code files   : Python, JavaScript, Shell, Rust, Assembly (read/syntax-highlight with pygments)
Execution    : .py → exec/subprocess, .js → node subprocess, .sh/.bash → sh subprocess

CLI commands (via NiblitCore._cmd_file)
----------------------------------------
file read <path>               — Read and display a file
file write <path> <content>    — Write content to a file (creates / overwrites)
file edit <path> <old>==><new> — Replace text inside a file
file execute <path> [args...]  — Execute a script/file
file detect <path>             — Detect file type and best handler
file formats                   — List all registered format handlers
file status                    — Handler registry summary
"""

from __future__ import annotations

import csv
import io
import json
import logging
import mimetypes
import os
import subprocess
import sys
import zipfile
import tarfile
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

log = logging.getLogger("UniversalFileManager")

# ─────────────────────────────────────────────────────────────────────────────
# Optional library detection (all soft-deps)
# ─────────────────────────────────────────────────────────────────────────────

def _try_import(name: str) -> Optional[Any]:
    try:
        import importlib
        return importlib.import_module(name)
    except Exception:
        return None


_yaml    = _try_import("yaml")
_tomllib = _try_import("tomllib") or _try_import("tomli")
_PIL     = _try_import("PIL.Image")
_pdfp    = _try_import("pdfplumber")
_pypdf   = _try_import("PyPDF2") if not _pdfp else None
_docx    = _try_import("docx")
_openpyxl= _try_import("openpyxl")
_sf      = _try_import("soundfile")
_pycdlib = _try_import("pycdlib")
_pygments_lex  = _try_import("pygments.lexers")
_pygments_fmt  = _try_import("pygments.formatters")
_pygments_high = _try_import("pygments.highlight")


# ─────────────────────────────────────────────────────────────────────────────
# Handler registry
# ─────────────────────────────────────────────────────────────────────────────

class FileHandler:
    """Abstract base for a file-type handler."""

    # Human-readable list of extensions this handler covers
    extensions: Tuple[str, ...] = ()
    description: str = ""

    def can_handle(self, path: Path) -> bool:
        return path.suffix.lower() in self.extensions

    def read(self, path: Path) -> str:
        raise NotImplementedError

    def write(self, path: Path, content: str) -> str:
        raise NotImplementedError

    def edit(self, path: Path, old: str, new: str) -> str:
        """Replace the first occurrence of *old* with *new* in the file."""
        text = self.read(path)
        if old not in text:
            return f"❌ '{old[:60]}...' not found in {path.name}"
        updated = text.replace(old, new, 1)
        return self.write(path, updated)

    def execute(self, path: Path, args: List[str]) -> str:
        return f"❌ Execution not supported for {path.suffix} files"

    @property
    def available(self) -> bool:
        return True  # override if depends on optional lib


# ── Plain-text handler ────────────────────────────────────────────────────────

class TextHandler(FileHandler):
    extensions = (".txt", ".md", ".rst", ".log", ".cfg", ".ini", ".env", ".toml")
    description = "Plain text / config files"

    def read(self, path: Path) -> str:
        return path.read_text(encoding="utf-8", errors="replace")

    def write(self, path: Path, content: str) -> str:
        path.write_text(content, encoding="utf-8")
        return f"✅ Written {len(content)} chars to {path}"


# ── JSON / JSONL handler ──────────────────────────────────────────────────────

class JsonHandler(FileHandler):
    extensions = (".json", ".jsonl")
    description = "JSON / JSON-Lines"

    def read(self, path: Path) -> str:
        raw = path.read_text(encoding="utf-8", errors="replace")
        try:
            if path.suffix == ".jsonl":
                objs = [json.loads(line) for line in raw.splitlines() if line.strip()]
                return json.dumps(objs[:20], indent=2) + (
                    f"\n… ({len(objs)} total lines)" if len(objs) > 20 else ""
                )
            obj = json.loads(raw)
            dumped = json.dumps(obj, indent=2)
            return dumped[:4000] + ("\n… (truncated)" if len(dumped) > 4000 else "")
        except Exception as exc:
            return f"[JSON parse error: {exc}]\n{raw[:2000]}"

    def write(self, path: Path, content: str) -> str:
        # Validate before writing
        try:
            json.loads(content)
        except Exception as exc:
            return f"❌ Invalid JSON: {exc}"
        path.write_text(content, encoding="utf-8")
        return f"✅ JSON written to {path}"


# ── CSV handler ───────────────────────────────────────────────────────────────

class CsvHandler(FileHandler):
    extensions = (".csv", ".tsv")
    description = "CSV / TSV spreadsheets"

    def read(self, path: Path) -> str:
        delim = "\t" if path.suffix == ".tsv" else ","
        rows: List[List[str]] = []
        try:
            with path.open(newline="", encoding="utf-8", errors="replace") as fh:
                reader = csv.reader(fh, delimiter=delim)
                for i, row in enumerate(reader):
                    if i >= 50:
                        rows.append(["… (truncated)"])
                        break
                    rows.append(row)
            return "\n".join(" | ".join(r) for r in rows)
        except Exception as exc:
            return f"❌ CSV read error: {exc}"

    def write(self, path: Path, content: str) -> str:
        path.write_text(content, encoding="utf-8")
        return f"✅ CSV written to {path}"


# ── YAML handler ──────────────────────────────────────────────────────────────

class YamlHandler(FileHandler):
    extensions = (".yaml", ".yml")
    description = "YAML"

    @property
    def available(self) -> bool:
        return _yaml is not None

    def read(self, path: Path) -> str:
        if not _yaml:
            return path.read_text(encoding="utf-8", errors="replace")
        try:
            data = _yaml.safe_load(path.read_text(encoding="utf-8", errors="replace"))
            return json.dumps(data, indent=2, default=str)[:4000]
        except Exception as exc:
            return f"[YAML parse error: {exc}]"

    def write(self, path: Path, content: str) -> str:
        path.write_text(content, encoding="utf-8")
        return f"✅ YAML written to {path}"


# ── PDF handler ───────────────────────────────────────────────────────────────

class PdfHandler(FileHandler):
    extensions = (".pdf",)
    description = "PDF documents"

    @property
    def available(self) -> bool:
        return bool(_pdfp or _pypdf)

    def read(self, path: Path) -> str:
        if _pdfp:
            try:
                import pdfplumber
                with pdfplumber.open(str(path)) as pdf:
                    text = "\n".join(page.extract_text() or "" for page in pdf.pages[:10])
                return text[:5000] or "[No extractable text]"
            except Exception as exc:
                return f"❌ pdfplumber error: {exc}"
        if _pypdf:
            try:
                import PyPDF2
                with open(str(path), "rb") as fh:
                    reader = PyPDF2.PdfReader(fh)
                    text = "\n".join(
                        page.extract_text() or "" for page in reader.pages[:10]
                    )
                return text[:5000] or "[No extractable text]"
            except Exception as exc:
                return f"❌ PyPDF2 error: {exc}"
        return "❌ PDF reading requires pdfplumber or PyPDF2 (pip install pdfplumber)"

    def write(self, path: Path, content: str) -> str:
        return "❌ Direct PDF writing not supported (use a template library)"


# ── DOCX handler ──────────────────────────────────────────────────────────────

class DocxHandler(FileHandler):
    extensions = (".docx",)
    description = "Microsoft Word documents"

    @property
    def available(self) -> bool:
        return _docx is not None

    def read(self, path: Path) -> str:
        if not _docx:
            return "❌ python-docx not installed (pip install python-docx)"
        try:
            import docx as _d
            doc = _d.Document(str(path))
            return "\n".join(p.text for p in doc.paragraphs)[:5000]
        except Exception as exc:
            return f"❌ DOCX read error: {exc}"

    def write(self, path: Path, content: str) -> str:
        if not _docx:
            return "❌ python-docx not installed"
        try:
            import docx as _d
            doc = _d.Document()
            for para in content.split("\n"):
                doc.add_paragraph(para)
            doc.save(str(path))
            return f"✅ DOCX written to {path}"
        except Exception as exc:
            return f"❌ DOCX write error: {exc}"


# ── Excel handler ─────────────────────────────────────────────────────────────

class ExcelHandler(FileHandler):
    extensions = (".xlsx", ".xls", ".ods")
    description = "Excel / OpenDocument spreadsheets"

    @property
    def available(self) -> bool:
        return _openpyxl is not None

    def read(self, path: Path) -> str:
        if not _openpyxl:
            return "❌ openpyxl not installed (pip install openpyxl)"
        try:
            import openpyxl
            wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
            lines: List[str] = []
            for sheet in wb.sheetnames[:3]:
                ws = wb[sheet]
                lines.append(f"=== Sheet: {sheet} ===")
                for i, row in enumerate(ws.iter_rows(values_only=True)):
                    if i >= 50:
                        lines.append("… (truncated)")
                        break
                    lines.append(" | ".join(str(c) for c in row))
            return "\n".join(lines)[:5000]
        except Exception as exc:
            return f"❌ Excel read error: {exc}"

    def write(self, path: Path, content: str) -> str:
        return "❌ Direct Excel writing: supply CSV and convert, or use openpyxl directly"


# ── Image handler ─────────────────────────────────────────────────────────────

class ImageHandler(FileHandler):
    extensions = (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp")
    description = "Images (Pillow)"

    @property
    def available(self) -> bool:
        return _PIL is not None

    def read(self, path: Path) -> str:
        if not _PIL:
            return "❌ Pillow not installed (pip install Pillow)"
        try:
            from PIL import Image
            img = Image.open(str(path))
            return (
                f"🖼 Image: {path.name}\n"
                f"  Format: {img.format}\n"
                f"  Size:   {img.size[0]} × {img.size[1]} px\n"
                f"  Mode:   {img.mode}"
            )
        except Exception as exc:
            return f"❌ Image read error: {exc}"

    def write(self, path: Path, content: str) -> str:
        return "❌ Image creation from text not supported — use Pillow directly"


# ── Audio handler ─────────────────────────────────────────────────────────────

class AudioHandler(FileHandler):
    extensions = (".wav", ".flac", ".ogg", ".mp3", ".aiff")
    description = "Audio files (soundfile)"

    @property
    def available(self) -> bool:
        return _sf is not None

    def read(self, path: Path) -> str:
        if not _sf:
            return "❌ soundfile not installed (pip install soundfile)"
        try:
            import soundfile as sf
            info = sf.info(str(path))
            return (
                f"🔊 Audio: {path.name}\n"
                f"  Format:    {info.format}\n"
                f"  Subtype:   {info.subtype}\n"
                f"  Duration:  {info.duration:.2f}s\n"
                f"  Samplerate:{info.samplerate} Hz\n"
                f"  Channels:  {info.channels}"
            )
        except Exception as exc:
            return f"❌ Audio read error: {exc}"

    def write(self, path: Path, content: str) -> str:
        return "❌ Audio writing from text not supported"


# ── Archive handler ───────────────────────────────────────────────────────────

class ArchiveHandler(FileHandler):
    extensions = (".zip", ".tar", ".gz", ".bz2", ".xz", ".tar.gz", ".tgz")
    description = "Archives (zip / tar)"

    def read(self, path: Path) -> str:
        try:
            if zipfile.is_zipfile(str(path)):
                with zipfile.ZipFile(str(path)) as zf:
                    names = zf.namelist()
                    return (
                        f"📦 ZIP archive ({len(names)} entries):\n"
                        + "\n".join(f"  {n}" for n in names[:40])
                        + ("\n  …" if len(names) > 40 else "")
                    )
            if tarfile.is_tarfile(str(path)):
                with tarfile.open(str(path)) as tf:
                    names = tf.getnames()
                    return (
                        f"📦 TAR archive ({len(names)} entries):\n"
                        + "\n".join(f"  {n}" for n in names[:40])
                        + ("\n  …" if len(names) > 40 else "")
                    )
            return "❌ Not a recognised archive format"
        except Exception as exc:
            return f"❌ Archive read error: {exc}"

    def write(self, path: Path, content: str) -> str:
        return "❌ Writing archives from text not supported"


# ── ISO / disk image handler ──────────────────────────────────────────────────

class IsoHandler(FileHandler):
    extensions = (".iso",)
    description = "ISO disk images (pycdlib)"

    @property
    def available(self) -> bool:
        return _pycdlib is not None

    def read(self, path: Path) -> str:
        if not _pycdlib:
            return "❌ pycdlib not installed (pip install pycdlib)"
        try:
            import pycdlib
            iso = pycdlib.PyCdlib()
            iso.open(str(path))
            children: List[str] = []
            for child in iso.listchildren("/"):
                children.append(child.file_identifier.decode("utf-8", errors="replace"))
            iso.close()
            return (
                f"💿 ISO image: {path.name} ({len(children)} root entries)\n"
                + "\n".join(f"  {c}" for c in children[:40])
            )
        except Exception as exc:
            return f"❌ ISO read error: {exc}"

    def write(self, path: Path, content: str) -> str:
        return "❌ Writing ISO images not supported"


# ── Code handler (syntax-highlighted read + execution) ───────────────────────

class CodeHandler(FileHandler):
    extensions = (
        ".py", ".js", ".ts", ".sh", ".bash", ".rs", ".asm", ".s",
        ".c", ".cpp", ".go", ".java", ".rb", ".php",
    )
    description = "Source code (read + execute)"

    def read(self, path: Path) -> str:
        text = path.read_text(encoding="utf-8", errors="replace")
        # Attempt syntax highlighting (stripped of ANSI for plain output)
        if _pygments_lex and _pygments_fmt and _pygments_high:
            try:
                from pygments import highlight
                from pygments.lexers import get_lexer_for_filename
                from pygments.formatters import TerminalFormatter
                lexer = get_lexer_for_filename(str(path))
                return highlight(text, lexer, TerminalFormatter())[:8000]
            except Exception:
                pass
        return text[:8000]

    def write(self, path: Path, content: str) -> str:
        path.write_text(content, encoding="utf-8")
        return f"✅ Source file written to {path}"

    def execute(self, path: Path, args: List[str]) -> str:  # noqa: C901
        ext = path.suffix.lower()
        arg_str = " ".join(args)

        def _run(cmd: List[str]) -> str:
            try:
                result = subprocess.run(
                    cmd,
                    capture_output=True,
                    text=True,
                    timeout=30,
                )
                out = (result.stdout or "")[:3000]
                err = (result.stderr or "")[:1000]
                rc = result.returncode
                lines = []
                if out:
                    lines.append(out)
                if err:
                    lines.append(f"[stderr]\n{err}")
                lines.append(f"[exit code: {rc}]")
                return "\n".join(lines)
            except subprocess.TimeoutExpired:
                return "❌ Execution timed out (30s limit)"
            except FileNotFoundError as exc:
                return f"❌ Interpreter not found: {exc}"
            except Exception as exc:
                return f"❌ Execution error: {exc}"

        if ext == ".py":
            cmd = [sys.executable, str(path)] + args
            return _run(cmd)
        if ext == ".js":
            return _run(["node", str(path)] + args)
        if ext in (".sh", ".bash"):
            return _run(["sh", str(path)] + args)
        return f"❌ Execution not configured for '{ext}' files"


# ─────────────────────────────────────────────────────────────────────────────
# Registry
# ─────────────────────────────────────────────────────────────────────────────

_ALL_HANDLERS: List[FileHandler] = [
    JsonHandler(),
    CsvHandler(),
    YamlHandler(),
    PdfHandler(),
    DocxHandler(),
    ExcelHandler(),
    ImageHandler(),
    AudioHandler(),
    ArchiveHandler(),
    IsoHandler(),
    CodeHandler(),
    TextHandler(),  # fallback — must be last
]


def _get_handler(path: Path) -> FileHandler:
    for h in _ALL_HANDLERS:
        if h.can_handle(path):
            return h
    return TextHandler()  # ultimate fallback


# ─────────────────────────────────────────────────────────────────────────────
# Main manager class
# ─────────────────────────────────────────────────────────────────────────────

class UniversalFileManager:
    """
    Niblit's universal file I/O, edit, and execution layer.

    All methods return human-readable strings so they can be surfaced
    directly as CLI / chat responses.
    """

    # ── Read ──────────────────────────────────────────────────────────────────

    def read_file(self, path_str: str) -> str:
        path = Path(path_str).expanduser()
        if not path.exists():
            return f"❌ File not found: {path}"
        handler = _get_handler(path)
        log.info("read_file: %s → %s", path, type(handler).__name__)
        try:
            return handler.read(path)
        except Exception as exc:
            return f"❌ Read error: {exc}"

    # ── Write ─────────────────────────────────────────────────────────────────

    def write_file(self, path_str: str, content: str) -> str:
        path = Path(path_str).expanduser()
        path.parent.mkdir(parents=True, exist_ok=True)
        handler = _get_handler(path)
        log.info("write_file: %s → %s", path, type(handler).__name__)
        try:
            return handler.write(path, content)
        except Exception as exc:
            return f"❌ Write error: {exc}"

    # ── Edit ──────────────────────────────────────────────────────────────────

    def edit_file(self, path_str: str, old: str, new: str) -> str:
        path = Path(path_str).expanduser()
        if not path.exists():
            return f"❌ File not found: {path}"
        handler = _get_handler(path)
        log.info("edit_file: %s → %s", path, type(handler).__name__)
        try:
            return handler.edit(path, old, new)
        except Exception as exc:
            return f"❌ Edit error: {exc}"

    # ── Execute ───────────────────────────────────────────────────────────────

    def execute_file(self, path_str: str, args: Optional[List[str]] = None) -> str:
        path = Path(path_str).expanduser()
        if not path.exists():
            return f"❌ File not found: {path}"
        handler = _get_handler(path)
        log.info("execute_file: %s → %s", path, type(handler).__name__)
        try:
            return handler.execute(path, args or [])
        except Exception as exc:
            return f"❌ Execute error: {exc}"

    def write_temp(self, content: str, suffix: str = ".txt") -> Dict[str, str]:
        """Write *content* to a temporary file and return its details.

        Uses ``io``, ``os``, and ``tempfile`` to create a named temporary file
        that is not automatically deleted on close.  This lets other Niblit
        components (code compiler, evolve engine, etc.) pass large content
        through the file system without polluting the working directory.

        Returns a dict with keys ``path``, ``size``, ``suffix``.
        """
        fd, tmp_path = tempfile.mkstemp(suffix=suffix)
        try:
            with io.open(fd, "w", encoding="utf-8") as fh:
                fh.write(content)
        except Exception:
            try:
                os.close(fd)
            except Exception:
                pass
            raise
        return {
            "path": tmp_path,
            "size": os.path.getsize(tmp_path),
            "suffix": suffix,
        }

    # ── Detect ────────────────────────────────────────────────────────────────

    def detect(self, path_str: str) -> str:
        path = Path(path_str).expanduser()
        mime, _ = mimetypes.guess_type(str(path))
        handler = _get_handler(path)
        return (
            f"🔍 {path.name}\n"
            f"  Extension : {path.suffix or '(none)'}\n"
            f"  MIME type : {mime or 'unknown'}\n"
            f"  Handler   : {type(handler).__name__} — {handler.description}\n"
            f"  Available : {'✅' if handler.available else '❌ (missing dependency)'}"
        )

    # ── List formats ─────────────────────────────────────────────────────────

    def list_formats(self) -> str:
        lines = ["Registered file format handlers:"]
        for h in _ALL_HANDLERS:
            exts = ", ".join(h.extensions) if h.extensions else "(fallback)"
            avail = "✅" if h.available else "❌"
            lines.append(f"  {avail} {type(h).__name__:<20} {exts:<40} {h.description}")
        return "\n".join(lines)

    # ── Status ────────────────────────────────────────────────────────────────

    def status(self) -> str:
        available = sum(1 for h in _ALL_HANDLERS if h.available)
        return (
            f"📁 UniversalFileManager | handlers={len(_ALL_HANDLERS)} "
            f"({available} available) | "
            f"pdf={'✅' if _pdfp or _pypdf else '❌'} "
            f"docx={'✅' if _docx else '❌'} "
            f"xlsx={'✅' if _openpyxl else '❌'} "
            f"image={'✅' if _PIL else '❌'} "
            f"audio={'✅' if _sf else '❌'} "
            f"iso={'✅' if _pycdlib else '❌'}"
        )


# ── Singleton ─────────────────────────────────────────────────────────────────
_manager_instance: Optional[UniversalFileManager] = None


def get_file_manager() -> UniversalFileManager:
    """Return the process-level UniversalFileManager singleton."""
    global _manager_instance  # pylint: disable=global-statement
    if _manager_instance is None:
        _manager_instance = UniversalFileManager()
    return _manager_instance
